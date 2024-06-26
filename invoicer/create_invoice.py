import logging
from datetime import datetime, timedelta

from dateutil import relativedelta
from docx.api import Document
from docx.opc.exceptions import PackageNotFoundError
from docx2pdf import convert

CONTRACTOR_DATA = "Jdt John Doe\nNIP:0123456789\nKonto: Test bank\n31 0000 0000 1111 0000 0000 0000"
CONTRACTOR_NAME = "John Doe"
CLIENT_DATA = "Very long name for contractor company\ntest street 15/4 floor, 00-1111 Test City\nNIP: 111-000-00-00"


def fill_invoice(price_net, value_net, value_gross, tax, tax_value, salary_in_words, save_invoice_path):
    try:
        document = Document("../invoicer/data_source/inv-template.docx")

        text_data_dict = inv_text_data(value_gross, salary_in_words)
        table_data_dict = inv_table_data(price_net, value_net, tax, tax_value, value_gross)

        for paragraph in document.paragraphs:

            for key, value in text_data_dict.items():
                if key in paragraph.text:
                    for run in paragraph.runs:
                        run.text = run.text.replace(key, str(value))

        for table in document.tables:
            for row in table.rows:
                for key, value in table_data_dict.items():
                    for cell in row.cells:
                        if key in cell.text:
                            for run in cell.paragraphs[0].runs:
                                run.text = run.text.replace(key, str(value))

        document.save(f"{save_invoice_path}.docx")
        convert(f"{save_invoice_path}.docx")

    except PackageNotFoundError as e:
        logging.exception(e)


def inv_text_data(value_gross, salary_in_words):
    text_data = {
        "[Invoice Number]": get_invoice_number(),
        "[Date Of Completion]": get_date_of_completion(),
        "[Contractor Data]": CONTRACTOR_DATA,
        "[Client Data]": CLIENT_DATA,
        "[Pay Day]": get_day_of_payment(),
        "[Value Gross]": value_gross,
        "[Salary in Words]": salary_in_words
    }
    return text_data


def inv_table_data(price_net, value_net, tax, tax_value, value_gross):
    table_data = {
        "[Price Net]": f"{price_net:.2f}",
        "[Value Net]": f"{value_net:.2f}",
        "[Tax]": f"{tax}%",
        "[Tax Value]": tax_value,
        "[Value Gross]": value_gross
    }
    return table_data


def get_date_of_completion():
    next_month = datetime.today().replace(day=28) + timedelta(days=4)
    last_day_of_month = next_month - timedelta(days=next_month.day)
    return last_day_of_month.date().strftime("%d-%m-%Y")


def get_invoice_number():
    return datetime.today().date().strftime("%m/%y")


def get_day_of_payment():
    day_of_payment = datetime.today().replace(day=21) + relativedelta.relativedelta(months=1)
    return day_of_payment.strftime("%d.%m.%Y")


def get_invoice_title():
    invoice_title_date = datetime.today().date().strftime("%m.%y")
    return f"{CONTRACTOR_NAME} -  Faktura Vat {invoice_title_date}"
